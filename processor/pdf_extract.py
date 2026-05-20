"""Download and extract text from RFP attachment PDFs.

Used to backfill the ``description`` field for SAM.gov records whose
inline ``noticedesc`` body is empty — for RFIs / Sources Sought / Special
Notices the real scope lives only in attached PDFs.

Strategy:
    1. Pick the most likely "spec" document from ``metadata.documents`` —
       prefer ``type == "primary_spec"`` and labels containing RFP/RFQ/SOW/
       Statement of Work / Scope; skip clearly-administrative attachments
       (response spreadsheets, Q&A, bidder conference materials).
    2. Download the PDF via the document's ``source_url`` (the original
       SAM.gov direct download — Drive URLs aren't publicly accessible
       without OAuth). Skip non-PDFs (.xlsx, .docx, .zip).
    3. Cache the extracted text under ``scraper/cache/samgov/pdf_text/<id>.txt``
       so subsequent runs reuse it instead of refetching.
"""

from __future__ import annotations

import io
import os
import re
from pathlib import Path
from typing import Any, Iterable, Optional

_PROJECT_ROOT = Path(__file__).resolve().parent.parent
_DEFAULT_CACHE_DIR = _PROJECT_ROOT / "scraper" / "cache" / "samgov" / "pdf_text"

# Attachment labels that aren't the actual spec.
_DOC_LABEL_BLOCKLIST = (
    "response_spreadsheet",
    "response spreadsheet",
    "vendor question",
    "q&a",
    "q_and_a",
    "questions",
    "bidder",
    "conference",
    "attendance",
    "presentation",
    "sign-in",
    "addendum",
    "amendment",
    "errata",
)

_DOC_LABEL_SPEC_HINTS = (
    "rfp",
    "rfq",
    "rfi",
    "sow",
    "scope",
    "spec",
    "specification",
    "exhibit",
    "statement of work",
    "requirement",
    "srd",  # SAM uses "SRD" for system requirements documents
    "pws",  # performance work statement
    "soo",  # statement of objectives
)


def _safe_cache_key(notice_id: str) -> str:
    return re.sub(r"[^\w.\-]+", "_", notice_id or "").strip("._") or "notice"


def _doc_kind(label: str, url: str) -> Optional[str]:
    """Return 'pdf' or 'docx' if the attachment is a readable text document,
    else ``None`` (xlsx, zip, etc. are unreadable here)."""
    lbl = (label or "").lower()
    if lbl.endswith(".pdf"):
        return "pdf"
    if lbl.endswith(".docx"):
        return "docx"
    if any(lbl.endswith(ext) for ext in (".xlsx", ".xls", ".zip", ".pptx", ".csv", ".doc")):
        return None
    # No extension hint — assume PDF (the SAM source_url returns octet-stream
    # for almost everything; we sniff the magic bytes after download).
    return "pdf"


def _rank_documents(documents: list[dict[str, Any]]) -> list[tuple[str, dict[str, Any]]]:
    """Return ``(kind, doc)`` tuples in best-spec-first order, dropping
    obviously admin attachments and unreadable formats."""
    candidates: list[tuple[int, str, dict[str, Any]]] = []
    for doc in documents or []:
        if not isinstance(doc, dict):
            continue
        label = (doc.get("label") or "").strip()
        url = doc.get("source_url") or doc.get("url") or ""
        if not url:
            continue
        kind = _doc_kind(label, url)
        if not kind:
            continue
        low = label.lower()
        if any(bad in low for bad in _DOC_LABEL_BLOCKLIST):
            continue
        score = 0
        if doc.get("type") == "primary_spec":
            score -= 10  # most preferred (sort ascending)
        if any(hint in low for hint in _DOC_LABEL_SPEC_HINTS):
            score -= 5
        # Larger files tend to be the real spec, not a cover letter.
        size = doc.get("size_bytes") or 0
        if isinstance(size, int):
            score -= min(size // 50000, 4)  # cap the bonus at -4
        # Prefer PDFs over docx when both exist for the same RFP — PDFs in
        # SAM are typically the canonical SOW; docx are draft/work versions.
        if kind == "docx":
            score += 1
        candidates.append((score, kind, doc))
    candidates.sort(key=lambda t: t[0])
    return [(kind, doc) for _, kind, doc in candidates]


def _download_bytes(url: str, *, timeout: int = 60) -> Optional[bytes]:
    """Download a URL; return bytes only if it looks like a real document."""
    import requests  # type: ignore

    try:
        resp = requests.get(url, timeout=timeout, allow_redirects=True)
    except Exception:
        return None
    if resp.status_code != 200 or len(resp.content) < 256:
        return None
    return resp.content


def _extract_text_from_pdf_bytes(data: bytes, *, max_pages: int = 30) -> str:
    """Run pypdf over ``data`` and return cleaned text (first ``max_pages`` pages)."""
    if data[:5] != b"%PDF-":
        return ""
    try:
        import pypdf  # type: ignore
    except ImportError:
        return ""

    try:
        reader = pypdf.PdfReader(io.BytesIO(data), strict=False)
    except Exception:
        return ""

    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        if i >= max_pages:
            break
        try:
            txt = page.extract_text() or ""
        except Exception:
            continue
        if txt:
            parts.append(txt)
    return _clean_pdf_text("\n".join(parts))


def _extract_text_from_docx_bytes(data: bytes) -> str:
    """Parse a .docx (Office Open XML) via python-docx."""
    if data[:4] != b"PK\x03\x04":  # docx is a zip
        return ""
    try:
        import docx  # type: ignore  # python-docx
    except ImportError:
        return ""

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception:
        return ""

    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Also pull table cells, which is where SAM RFI forms put most content.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = (cell.text or "").strip()
                if txt:
                    paragraphs.append(txt)
    return _clean_pdf_text("\n".join(paragraphs))


# Common first-real-section headings used in RFP / SOW PDFs. We trim
# everything before the first one of these so cover pages, change history,
# and tables of contents don't drown out the actual scope text.
_SECTION_HEAD_RE = re.compile(
    r"^\s*(?:\d+(?:\.\d+)*\s+)?"  # optional outline numbering: "1 ", "1.1 "
    r"(?:scope|introduction|overview|background|purpose|"
    r"statement\s+of\s+work|sow|objective(?:s)?|"
    r"executive\s+summary|general\s+information|"
    r"summary|description\s+of\s+services)\s*$",
    re.IGNORECASE,
)

# Lines that are just table-of-contents entries — runs of dot leaders.
_TOC_DOT_LEADERS_RE = re.compile(r"\.{4,}")

# Common English stop words; used to detect natural-language prose.
_STOPWORDS = frozenset({
    "the", "of", "is", "a", "an", "to", "in", "for", "and", "or", "be",
    "with", "from", "on", "by", "as", "that", "this", "are", "will",
    "have", "has", "their", "they", "its", "it", "at", "which", "these",
    "those", "we", "our", "shall", "must",
})

# Patterns that mean a line is administrative / certification scaffolding,
# not project narrative. Used to drop noisy lines from the description.
_NON_NARRATIVE_PATTERNS_CI = (
    r"\bi\s+certify\b",
    r"\bper\s+vaar\b",
    r"\bper\s+far\b",
    r"\bname,\s*date\s*:",
    r"\bapproved\s+by\s*:",
    r"\bsignature\s*:_+",
    r"\bdistribution\s+statement\b",
    r"^\s*brand\s+name\s+item\s*:",
    r"^\s*essential/significant\s+physical",
    r"^\s*salient\s+characteristics\b",
    r"^\s*supervisory\s+",
    r"^\s*contracting\s+officer\b",
    r"^\s*vacchcs\b",
    r"^\s*version\s+\d",
    r"^\s*far\s+\d",
    r"^\s*\d{1,2}/\d{1,2}/\d{2,4}\s*$",
    # Numbered form fields with short labels: "2) Complete generic identification – ..."
    r"^\s*\d+\)\s*[a-z][a-z\-/ ]{2,50}(?:\s*[–—\-:])",
    # Numbered section headings: "5) Capabilities" / "6) Software Requirements".
    r"^\s*\d+\)\s*[a-z][a-z &/]{2,35}\s*$",
    # SAM PDF form-field label lines (uppercase phrase with optional trailing *).
    # e.g. "SUBJECT*", "POINT OF CONTACT*", "SOLICITATION NUMBER*".
    r"^\s*[A-Z][A-Z0-9'’\s,&/().\-]{1,80}\*\s*$",
)
_NON_NARRATIVE_CI_RE = re.compile("|".join(_NON_NARRATIVE_PATTERNS_CI), re.IGNORECASE | re.MULTILINE)

# Short title-cased section subheaders (e.g. "Streamlined Record Keeping",
# "Capabilities"). Case-sensitive on purpose so we don't drop real prose.
_TITLE_CASE_HEADER_RE = re.compile(r"^\s*(?:[A-Z][A-Za-z]+\s*){1,4}$")

# ALL-CAPS banner lines (e.g. "NOTICE OF INTENT TO SOLE SOURCE",
# "GENERAL INFORMATION", "DESCRIPTION").
_ALLCAPS_BANNER_RE = re.compile(r"^\s*[A-Z][A-Z0-9\s,&/().\-]{4,80}\s*$")

# Pure numeric / code lines: zip codes, solicitation numbers, phone numbers.
# e.g. "01730", "36C24126Q0465", "541511", "5-22-70-25-001 to 359".
_CODE_LINE_RE = re.compile(r"^\s*[A-Z0-9][A-Z0-9\-\s/]{1,40}\s*$")


def _looks_like_non_narrative_line(stripped: str) -> bool:
    if _NON_NARRATIVE_CI_RE.search(stripped):
        return True
    if _TITLE_CASE_HEADER_RE.match(stripped):
        return True
    if _ALLCAPS_BANNER_RE.match(stripped):
        return True
    # Pure number / code line with no lowercase letters and very few stop words.
    if _CODE_LINE_RE.match(stripped):
        # Confirm it has no lowercase letters and is short — avoid eating
        # short legitimate sentences.
        if not any(c.islower() for c in stripped) and len(stripped) <= 50:
            return True
    return False


# Admin lead-in sentences that often start SAM notice bodies — we don't want
# them at the very top of the description. Same set used by enrich._FILLER_LEAD_PATTERNS
# but kept here to avoid a circular import.
_ADMIN_LEAD_PATTERNS = (
    "this is a combined synopsis",
    "this announcement constitutes",
    "this is a sources sought",
    "this is a request for information",
    "this notice does not constitute",
    "issued solely for information",
    "issued for market research",
    "is conducting market research",
    "is issuing this notice",
    "notice of intent",
    "offers are being requested",
    "see the attached",
    "please see the attached",
)


def _drop_leading_admin_sentences(text: str) -> str:
    """Strip leading sentences that are pure procurement-process boilerplate."""
    if not text:
        return text
    paragraphs = re.split(r"\n\s*\n", text)
    out_paragraphs: list[str] = []
    consumed_any = False
    for idx, para in enumerate(paragraphs):
        sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z“\"'])", para)
        kept_sentences: list[str] = []
        for s in sentences:
            low = s.lower().strip()
            if not consumed_any and any(p in low for p in _ADMIN_LEAD_PATTERNS):
                # Drop only leading admin sentences; once we keep one substantive
                # sentence, stop filtering (later admin context might be project-relevant).
                continue
            kept_sentences.append(s)
            consumed_any = True
        joined = " ".join(s.strip() for s in kept_sentences if s.strip()).strip()
        if joined:
            out_paragraphs.append(joined)
        elif out_paragraphs:
            # preserve paragraph break placeholder
            pass
    return "\n\n".join(out_paragraphs).strip()


def _is_narrative_paragraph(para: str) -> bool:
    """True when ``para`` reads like English prose (not a form field / header)."""
    stripped = para.strip()
    if len(stripped) < 60:
        return False
    if _looks_like_non_narrative_line(stripped):
        return False
    # Mostly-uppercase paragraphs are section banners or all-caps headers.
    letters = [c for c in stripped if c.isalpha()]
    if letters and sum(1 for c in letters if c.isupper()) / len(letters) > 0.55:
        return False
    # Need real English structure — at least a few stop words.
    tokens = re.findall(r"[A-Za-z]+", stripped.lower())
    if len(tokens) < 12:
        return False
    stop_hits = sum(1 for t in tokens if t in _STOPWORDS)
    if stop_hits < 3:
        return False
    return True


def _build_narrative_description(text: str, *, target_chars: int = 1800) -> str:
    """Reduce messy PDF text down to a few clean narrative paragraphs.

    Used in place of the raw PDF dump so SAM ``description`` fields read like
    a coherent project summary instead of a transcription of cover pages,
    form fields, and certifications.
    """
    if not text:
        return ""
    # First pass: drop lines that are obviously non-narrative scaffolding
    # (form fields, signatures, certifications, short title-case headers).
    cleaned_lines: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append("")
            continue
        if _looks_like_non_narrative_line(stripped):
            continue
        cleaned_lines.append(stripped)
    pruned = "\n".join(cleaned_lines)

    # Second pass: pick narrative paragraphs in document order until we have
    # enough content for a useful description.
    paragraphs = re.split(r"\n\s*\n", pruned)
    kept: list[str] = []
    chars = 0
    for para in paragraphs:
        para = re.sub(r"\s+\n", "\n", para).strip()
        if not _is_narrative_paragraph(para):
            continue
        kept.append(para)
        chars += len(para) + 2
        if chars >= target_chars:
            break
    return "\n\n".join(kept).strip()


# Anchor patterns to try in priority order. Inline anchors with colons are
# more specific than bare line headings (which sometimes match unrelated
# banners like "GENERAL INFORMATION" that just precede more boilerplate).
_HIGH_SIGNAL_ANCHOR_RE = re.compile(
    r"(?:^|[.!?]\s|\n)\s*"
    r"(?:\d+\.\s+)?"
    r"(?:statement\s+of\s+work|scope\s+of\s+work|description\s+of\s+services|"
    r"description|sow)\s*:\s*(?:\n|$)",
    re.IGNORECASE,
)


def _trim_to_section_body(text: str) -> str:
    """Find the most specific section anchor and return the text after it.

    Priority order:
        1. Inline "STATEMENT OF WORK:" / "DESCRIPTION:" anchors (with colon).
        2. Bare line-only section headings ("SCOPE", "INTRODUCTION", …).

    If neither fires, return the input unchanged.
    """
    if not text:
        return text

    # 1. High-signal anchor: take the LAST match (later sections in the doc
    # are typically the body, not the table of contents entry).
    matches = list(_HIGH_SIGNAL_ANCHOR_RE.finditer(text))
    if matches:
        return text[matches[-1].end() :]

    # 2. Fallback: bare section heading on its own line.
    lines = text.splitlines()
    for i, ln in enumerate(lines[:120]):
        if _SECTION_HEAD_RE.match(ln):
            return "\n".join(lines[i + 1 :])
    return text


def _clean_pdf_text(text: str) -> str:
    """Tidy PDF extractor output — strip distribution notices, blank pages, repeats."""
    if not text:
        return ""
    # Drop common SAM boilerplate / distribution notices that flood every page.
    boilerplate_patterns = (
        r"Distribution Statement A\.\s*Approved for public release[^\n]*",
        r"THIS PAGE LEFT INTENTIONALLY BLANK",
        r"^\s*Page \d+ of \d+\s*$",
        r"FOR OFFICIAL USE ONLY",
        r"CHANGE\s+HISTORY",
    )
    for pat in boilerplate_patterns:
        text = re.sub(pat, "", text, flags=re.IGNORECASE | re.MULTILINE)

    # Drop table-of-contents lines.
    lines = [ln for ln in text.splitlines() if not _TOC_DOT_LEADERS_RE.search(ln)]
    text = "\n".join(lines)

    # Trim away the cover page / TOC / boilerplate that precedes the real SOW.
    # Prefer high-signal "STATEMENT OF WORK:" / "DESCRIPTION:" anchors over
    # generic banners like "GENERAL INFORMATION"; fall back to bare section
    # headings if no specific anchor is present.
    text = _trim_to_section_body(text)

    # Normalize whitespace.
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    # Reduce to narrative paragraphs only — drops form scaffolding, signatures,
    # certifications. Falls back to the raw cleaned text if filtering left
    # nothing usable (rare; happens on heavily form-based docs).
    narrative = _build_narrative_description(text)
    narrative = _drop_leading_admin_sentences(narrative)
    return narrative or text


def extract_text_for_record(
    raw: dict[str, Any],
    *,
    cache_dir: os.PathLike[str] | str = _DEFAULT_CACHE_DIR,
    max_pages: int = 30,
    refresh: bool = False,
) -> str:
    """Return PDF-extracted scope text for ``raw``, or ``""`` if not available.

    Cached under ``scraper/cache/samgov/pdf_text/<external_id>.txt`` so
    re-runs don't re-download. Pass ``refresh=True`` to force re-extraction.
    """
    external_id = (raw.get("external_id") or "").strip()
    cache_path = Path(cache_dir) / f"{_safe_cache_key(external_id)}.txt"
    if not refresh and cache_path.is_file():
        return cache_path.read_text(encoding="utf-8")

    documents = ((raw.get("metadata") or {}).get("documents")) or []
    ranked = _rank_documents(documents)

    text = ""
    for kind, doc in ranked:
        url = doc.get("source_url") or doc.get("url")
        if not url:
            continue
        data = _download_bytes(url)
        if not data:
            continue
        if kind == "pdf":
            text = _extract_text_from_pdf_bytes(data, max_pages=max_pages)
        elif kind == "docx":
            text = _extract_text_from_docx_bytes(data)
        else:
            text = ""
        if text and len(text) >= 200:
            break

    cache_path.parent.mkdir(parents=True, exist_ok=True)
    cache_path.write_text(text or "", encoding="utf-8")
    return text or ""


def maybe_backfill_description(raw: dict[str, Any], **kwargs: Any) -> bool:
    """Mutate ``raw`` in place: fill ``description`` from PDFs when empty.

    Returns ``True`` if the description was backfilled, ``False`` otherwise.
    Only fires for records whose ``description`` is empty / whitespace and
    that have at least one document URL — applies to every source, but in
    practice this only fires for SAM.gov RFIs / Special Notices today.
    """
    existing = (raw.get("description") or "").strip()
    if existing:
        return False

    text = extract_text_for_record(raw, **kwargs)
    if not text:
        return False

    raw["description"] = text
    return True
